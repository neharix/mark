import datetime
import random
from io import BytesIO

import numpy as np
import pandas as pd
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import Group, User
from django.core.files.base import ContentFile
from django.db import utils
from django.http import HttpRequest, HttpResponse
from django.shortcuts import redirect, render
from django_ratelimit.decorators import ratelimit
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm, Pt, RGBColor
from encrypted_files.base import EncryptedFile

from .containers import *
from .models import *


@ratelimit(key="ip", rate="5/s")
def login_view(request):
    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["password"]
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect("home")
        else:
            return render(
                request,
                "views/login.html",
                {"message": "Ulanyjy adyňyz ýa-da/we açar sözüňiz nädogry"},
            )
    else:
        if request.user.is_authenticated:
            return redirect("home")
        else:
            return render(
                request,
                "views/login.html",
            )


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def logout_view(request):
    logout(request)
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def main(request: HttpRequest):
    if request.user.groups.contains(Group.objects.get(name="Moderator")):
        moderators_count = User.objects.filter(groups__name="Moderator").count()
        juries_count = User.objects.filter(groups__name="Jury").count()
        spectators_count = User.objects.filter(groups__name="Spectator").count()
        projects_count = Project.objects.all().count()
        rated_projects_count = Project.rated_objects.all().count()
        unrated_projects_count = Project.unrated_objects.all().count()
        schedules = [
            ScheduleContainer(schedule)
            for schedule in Schedule.objects.filter(
                date__gte=datetime.date.today()
            ).order_by("-date")[:5]
        ]
        return render(
            request,
            "views/main/moderator.html",
            {
                "moderators_count": moderators_count,
                "juries_count": juries_count,
                "spectators_count": spectators_count,
                "projects_count": projects_count,
                "rated_projects_count": rated_projects_count,
                "unrated_projects_count": unrated_projects_count,
                "schedules": schedules,
            },
        )
    elif request.user.groups.contains(Group.objects.get(name="Jury")):
        today = datetime.datetime.now().astimezone(pytz.timezone("Asia/Ashgabat"))
        context = {"projects_count": 0, "juries_count": 0}
        try:
            schedule = Schedule.objects.get(
                date__year=today.year, date__month=today.month, date__day=today.day
            )
            projects = []
            for pk in json.loads(schedule.quene_json):
                projects.append(Project.objects.get(pk=pk))
            context["is_participate"] = schedule.juries.filter(
                username=request.user.username
            ).exists()
            context["today"] = today

            context["marks"] = [
                SecondaryMarkContainer(mark)
                for mark in Mark.objects.filter(jury=request.user).order_by("-date")
            ]

            context["projects_count"] = len(json.loads(schedule.quene_json))
            context["juries_count"] = schedule.juries.all().count()
            context["projects"] = projects
            context["juries"] = schedule.juries.all()
        except:
            pass
        return render(
            request,
            "views/main/jury.html",
            context,
        )
    elif request.user.groups.contains(Group.objects.get(name="Spectator")):
        rated_projects_count = Project.rated_objects.all().count()
        rated_projects = []
        today = datetime.datetime.today().astimezone(pytz.timezone("Asia/Ashgabat"))

        for project in Project.objects.all():
            if Mark.objects.filter(project=project).exists():
                rated_projects.append(ProjectMarkContainer(project))
        rated_projects.sort(key=lambda e: e.percent)
        rated_projects.reverse()
        ls_projects = [project for project in rated_projects]
        ls_projects.sort(key=lambda e: e.mark_date)
        unrated_projects_count = Project.unrated_objects.all().count()
        directions = [
            DirectionContainer(direction) for direction in Direction.objects.all()
        ]
        return render(
            request,
            "views/main/spectator.html",
            {
                "rated_projects_count": rated_projects_count,
                "rated_projects": rated_projects[:10],
                "unrated_projects_count": unrated_projects_count,
                "directions": directions,
                "ls_projects": ls_projects,
            },
        )
    return render(request, "views/main/default.html")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def schedules(request: HttpRequest):
    if request.user.groups.contains(Group.objects.get(name="Moderator")):
        today = datetime.date.today()
        schedules = [
            ScheduleContainer(schedule)
            for schedule in Schedule.objects.filter(date__gte=today).order_by("-date")
        ]
        return render(
            request,
            "views/schedules.html",
            {
                "schedules": schedules,
            },
        )
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def add_project(request: HttpRequest):
    if request.user.groups.contains(Group.objects.get(name="Moderator")):
        directions = Direction.objects.all()
        if request.method == "POST":
            try:
                direction = Direction.objects.get(pk=int(request.POST["direction"]))
            except:
                return render(
                    request,
                    "views/add_project.html",
                    {
                        "directions": directions,
                        "message": "Ugry dogry giriziň!",
                    },
                )
            try:
                project = Project.objects.create(
                    agency=request.POST["agency"],
                    direction=direction,
                    full_name_of_manager=request.POST["manager_full_name"],
                    description=request.POST["description"],
                )
                if request.POST.get("second_member_full_name", False):
                    if request.POST["second_member_full_name"] != "ýok":
                        project.full_name_of_second_participant = request.POST[
                            "second_member_full_name"
                        ]
                if request.POST.get("third_member_full_name", False):
                    if request.POST["third_member_full_name"] != "ýok":
                        project.full_name_of_third_participant = request.POST[
                            "third_member_full_name"
                        ]
                project.save()
            except:
                return render(
                    request,
                    "views/add_project.html",
                    {"message": "Ýalňyşlyk döredi!", "directions": directions},
                )
        return render(request, "views/add_project.html", {"directions": directions})
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def add_project_using_xlsx(request: HttpRequest):
    if request.user.groups.contains(Group.objects.get(name="Moderator")):
        if request.method == "POST":
            print(request.FILES.get("xlsx").size)
            dataframe = pd.read_excel(request.FILES["xlsx"], engine="openpyxl")
            for index in range(len(dataframe["Ýolbaşçynyň F.A.A"])):
                try:
                    direction = Direction.objects.get(name=dataframe["Ugry"][index])
                except:
                    direction = Direction.objects.create(name=dataframe["Ugry"][index])

                project = Project.objects.create(
                    agency=dataframe["Edaranyň ady"][index],
                    direction=direction,
                    full_name_of_manager=dataframe["Ýolbaşçynyň F.A.A"][index],
                    description=dataframe["Taslamanyň beýany"][index],
                )

                if (
                    not type(dataframe["Ikinji agzanyň F.A.A"][index]) == np.float64
                    or not type(dataframe["Üçünji agzanyň F.A.A"][index]) == np.nan
                ):
                    project.full_name_of_second_participant = dataframe[
                        "Ikinji agzanyň F.A.A"
                    ][index]

                if (
                    not type(dataframe["Üçünji agzanyň F.A.A"][index]) == np.float64
                    or not type(dataframe["Üçünji agzanyň F.A.A"][index]) == np.nan
                ):
                    project.full_name_of_third_participant = dataframe[
                        "Üçünji agzanyň F.A.A"
                    ][index]
                project.save()

        return render(request, "views/add_project_using_xlsx.html")
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def add_schedule(request: HttpRequest):
    if request.user.groups.contains(Group.objects.get(name="Moderator")):
        return render(request, "views/add_schedule.html")
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def edit_schedule(request: HttpRequest, schedule_pk: int):
    if request.user.groups.contains(Group.objects.get(name="Moderator")):
        schedule = Schedule.objects.get(pk=schedule_pk)
        juries = schedule.juries.all()
        project_pks = json.loads(schedule.quene_json)
        projects = []
        for project_pk in project_pks:
            projects.append(Project.objects.get(pk=project_pk))
        return render(
            request,
            "views/edit_schedule.html",
            {"schedule": schedule, "projects": projects, "juries": juries},
        )
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def delete_schedule(request: HttpRequest, schedule_pk: int):
    if request.user.groups.contains(Group.objects.get(name="Moderator")):
        Schedule.objects.get(pk=schedule_pk).delete()
        return redirect("schedules")
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def mark_form(request: HttpRequest):
    context = {}
    if request.user.groups.contains(Group.objects.get(name="Jury")):
        if request.method == "POST":
            rated_project = Project.objects.get(pk=int(request.POST["project-pk"]))
            if Mark.objects.filter(jury=request.user, project=rated_project).exists():
                return redirect("mark_form")
            mark = Mark.objects.create(
                jury=request.user,
                mark=(
                    int(request.POST["mark"]) if int(request.POST["mark"]) >= 10 else 10
                ),
                project=rated_project,
            )
            if request.POST[f"description"] != "":
                mark.description = request.POST[f"description"]
            mark.save()

            today = datetime.datetime.today()
            schedule = Schedule.objects.get(
                date__year=today.year, date__month=today.month, date__day=today.day
            )

            juries = schedule.juries.all()
            rate_count = 0
            for jury in juries:
                if Mark.objects.filter(jury=jury, project=rated_project).exists():
                    rate_count += 1
            if rate_count == juries.count():
                print(rate_count)
                rated_project.rated = True
                rated_project.save()
        today = datetime.datetime.today()
        try:
            schedule = Schedule.objects.get(
                date__year=today.year, date__month=today.month, date__day=today.day
            )
        except:
            return redirect("home")
        if not schedule.juries.filter(pk=request.user.pk):
            return redirect("home")

        for project_pk in json.loads(schedule.quene_json):
            project = Project.objects.get(pk=project_pk)
            if (
                not Mark.objects.filter(
                    project=project,
                    jury=request.user,
                ).exists()
                and not project.rated
            ):
                context["project"] = project
                break
        else:
            return redirect("home")
        return render(request, "views/mark_form.html", context)
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def export_to_docx(request: HttpRequest, direction: str = None):
    if request.user.groups.contains(Group.objects.get(name="Spectator")):
        if direction != None:
            if not direction.isdigit():
                return redirect("projects_list")
            elif Direction.objects.filter(pk=int(direction)).exists():
                direction = int(direction)
            else:
                return redirect("projects_list")
        sort_by_marks = lambda e: e.percent
        rated_projects = []
        if direction != None:
            for project in Project.objects.filter(
                direction=Direction.objects.get(pk=direction)
            ):
                if Mark.objects.filter(project=project).exists():
                    rated_projects.append(ProjectMarkContainer(project))
        else:
            for project in Project.objects.all():
                if Mark.objects.filter(project=project).exists():
                    rated_projects.append(ProjectMarkContainer(project))
        rated_projects.sort(key=sort_by_marks)
        rated_projects.reverse()
        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"

        picture_paragraph = document.add_paragraph()
        picture_run = picture_paragraph.add_run()
        picture_run.add_picture(
            str(settings.BASE_DIR) + "/main/static/img/ministry-logo.png",
            width=Mm(25),
            height=Mm(25),
        )
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head = document.add_heading(level=0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = head.add_run(f'"Sanly Çözgüt - 2024" netijeleri')

        run.font.name = "Times New Roman"
        run.font.bold = True

        table = document.add_table(rows=len(rated_projects) + 1, cols=6)
        table.style = "Table Grid"

        for index in range(6):
            cell = table.cell(0, index)
            paragraph = cell.paragraphs[0]
            if index == 0:
                run = paragraph.add_run("№")
            elif index == 1:
                run = paragraph.add_run("Taslama")
            elif index == 2:
                run = paragraph.add_run("Bahasy")
            elif index == 3:
                run = paragraph.add_run("Ýolbaşçysy")
            elif index == 4:
                run = paragraph.add_run("Edarasy")
            elif index == 5:
                run = paragraph.add_run("Ugry")
            run.font.bold = True

        row_id = 1

        for project in rated_projects:

            table.cell(row_id, 0).text = f"{row_id}"

            table.cell(row_id, 1).text = f"{project.name}"
            table.cell(row_id, 2).text = f"{project.percent}%"
            table.cell(row_id, 3).text = f"{project.manager}"
            table.cell(row_id, 4).text = f"{project.agency}"
            table.cell(row_id, 5).text = f"{project.direction}"
            row_id += 1

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        response["Content-Disposition"] = (
            f'attachment; filename="netije-{datetime.datetime.now().strftime('%d.%m.%Y')}.docx"'
        )

        return response

    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def export_to_xlsx(request: HttpRequest, direction: str = None):
    if request.user.groups.contains(Group.objects.get(name="Spectator")):
        if direction != None:
            if not direction.isdigit():
                return redirect("projects_list")
            elif Direction.objects.filter(pk=int(direction)).exists():
                direction = int(direction)
            else:
                return redirect("projects_list")
        rated_projects = []
        if direction != None:
            for project in Project.objects.filter(
                direction=Direction.objects.get(pk=direction)
            ):
                if Mark.objects.filter(project=project).exists():
                    rated_projects.append(ProjectMarkContainer(project))
        else:
            for project in Project.objects.all():
                if Mark.objects.filter(project=project).exists():
                    rated_projects.append(ProjectMarkContainer(project))
        rated_projects.sort(key=lambda e: e.percent)
        rated_projects.reverse()
        dataframe_dict = {
            "Taslama": [],
            "Bahasy": [],
            "Ýolbaşçysy": [],
            "Edarasy": [],
            "Ugry": [],
        }

        for project in rated_projects:
            dataframe_dict["Taslama"].append(project.name)
            dataframe_dict["Bahasy"].append(f"{project.percent}%")
            dataframe_dict["Ýolbaşçysy"].append(f"{project.manager}")
            dataframe_dict["Edarasy"].append(f"{project.agency}")
            dataframe_dict["Ugry"].append(f"{project.direction}")

        dataframe = pd.DataFrame(dataframe_dict)
        response = HttpResponse(
            content_type="application/xlsx",
        )
        response["Content-Disposition"] = f'attachment; filename="results.xlsx"'
        with pd.ExcelWriter(response) as writer:
            dataframe.to_excel(writer, sheet_name="sheet1")
        return response
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def projects_list(request: HttpRequest, direction="all"):
    if request.user.groups.contains(Group.objects.get(name="Spectator")):
        projects = []

        if direction != "all":
            if Direction.objects.filter(pk=int(direction)):
                direction_obj = Direction.objects.get(pk=int(direction))
                direction = int(direction)
                for project in Project.objects.filter(direction=direction_obj):
                    if Mark.objects.filter(project=project).exists():
                        projects.append(project)
                return render(
                    request,
                    "views/projects_list.html",
                    {
                        "projects": projects,
                        "directions": Direction.objects.all(),
                        "dir": direction,
                    },
                )
            else:
                return redirect("projects_list")
        for project in Project.objects.all():
            if Mark.objects.filter(project=project).exists():
                projects.append(project)
        return render(
            request,
            "views/projects_list.html",
            {
                "projects": projects,
                "directions": Direction.objects.all(),
                "dir": "all",
            },
        )
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def project_result(request: HttpRequest, project_pk: int):
    if request.user.groups.contains(Group.objects.get(name="Spectator")):
        project = Project.objects.get(pk=project_pk)
        marks = [
            MarkContainer(mark)
            for mark in Mark.objects.filter(project=project).order_by("-date")
        ]
        schedule = get_projects_schedule(project)
        unparticipated_juries = [
            UnparticipatedJuryContainer(jury)
            for jury in get_unparticipated_juries(project, schedule)
        ]
        unparticipated_juries_count = len(unparticipated_juries)
        project_mark_container = ProjectMarkContainer(project)
        return render(
            request,
            "views/project_result.html",
            {
                "project": project,
                "unparticipated_juries_count": unparticipated_juries_count,
                "marks": numerate_containers(unparticipated_juries + marks),
                "project_mark_container": project_mark_container,
            },
        )


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def add_user(request: HttpRequest):
    if request.user.groups.contains(Group.objects.get(name="Moderator")):
        groups = Group.objects.all()
        if request.method == "POST":
            try:
                group = Group.objects.get(pk=int(request.POST["role"]))
                username = (
                    request.POST["first_name"].lower()
                    + request.POST["last_name"].lower()
                )
                user = User.objects.create(
                    last_name=request.POST["last_name"],
                    first_name=request.POST["first_name"],
                    username=username,
                )
                password = username + "".join(
                    [str(random.randint(0, 9)) for i in range(4)]
                )
                user.set_password(password)
                profile = Profile.objects.get(user=user)
                profile.password = password
                profile.save()

                user.groups.add(group)
                user.save()
                return redirect("users")
            except utils.IntegrityError:
                return render(
                    request,
                    "views/add_user.html",
                    {
                        "groups": groups,
                        "message": f"Ulanyjy eýýäm girizildi",
                    },
                )
            except:
                return render(
                    request,
                    "views/add_user.html",
                    {
                        "groups": groups,
                        "message": f"Näsazlyk ýüze çykdy",
                    },
                )

        return render(request, "views/add_user.html", {"groups": groups})
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def users(request: HttpRequest):
    if request.user.groups.contains(Group.objects.get(name="Moderator")):
        juries = User.objects.filter(groups__name="Jury")
        spectators = User.objects.filter(groups__name="Spectator")
        moderators = User.objects.filter(groups__name="Moderator")
        return render(
            request,
            "views/users.html",
            {
                "juries": juries,
                "spectators": spectators,
                "moderators": moderators,
            },
        )
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def delete_user(request: HttpRequest, user_pk: int):
    if request.user.groups.contains(Group.objects.get(name="Moderator")):
        User.objects.get(pk=user_pk).delete()
        return redirect("users")
    return redirect("home")


@ratelimit(key="ip", rate="5/s")
@login_required(login_url="/login/")
def edit_mark(request: HttpRequest, mark_pk: int):
    if request.user.groups.contains(Group.objects.get(name="Jury")):
        mark = Mark.objects.get(pk=mark_pk)
        if request.method == "POST":
            mark.mark = (
                int(request.POST["mark"]) if int(request.POST["mark"]) >= 10 else 10
            )
            mark.description = (
                request.POST["description"]
                if request.POST.get("description", False)
                else None
            )
            mark.save()
            return redirect("home")

        today = datetime.datetime.now().astimezone(pytz.timezone("Asia/Ashgabat"))
        if (
            mark.date.astimezone(pytz.timezone("Asia/Ashgabat")).day == today.day
            and mark.date.astimezone(pytz.timezone("Asia/Ashgabat")).month
            == today.month
            and mark.date.astimezone(pytz.timezone("Asia/Ashgabat")).year == today.year
            and today.hour < 20
        ):
            return render(
                request,
                "views/mark_form.html",
                {"project": mark.project, "mark": mark},
            )
    return redirect("home")
